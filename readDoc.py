#! python 3
from docx import Document
from docx.shared import Inches
from docx.image.png import Png

document = Document()
p = document.add_paragraph('hello world.\t')
p.add_run('hello again')
document.add_picture('phuong.png', width = Inches(0.01))
document.save('helloworld.docx')

# def getText(fullname):
#     doc = docx.Document(fullname)
#     fullText = []
#     for para in doc.paragraphs:
#         fullText.append(para.text)
#     return '\n'.join(fullText)
#
# print(getText('helloworld.docx'))
